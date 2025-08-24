"""File processing utilities for FileInASnap MCP Server"""

import os
import mimetypes
from pathlib import Path
from typing import Dict, Any, Optional, List, Union, BinaryIO
import aiofiles
import hashlib
import logging

logger = logging.getLogger(__name__)

# Optional dependencies
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logger.warning("Pillow not available, image processing disabled")

try:
    import cv2
    CV2_AVAILABLE = True
except ImportError:
    CV2_AVAILABLE = False
    logger.warning("OpenCV not available, video processing disabled")

try:
    import imagehash
    IMAGEHASH_AVAILABLE = True
except ImportError:
    IMAGEHASH_AVAILABLE = False
    logger.warning("imagehash not available, perceptual hashing disabled")

try:
    import PyPDF2
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    logger.warning("PyPDF2 not available, PDF processing disabled")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available, Word document processing disabled")

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.warning("openpyxl not available, Excel processing disabled")


class FileProcessor:
    """Handles file processing and metadata extraction"""
    
    def __init__(self):
        # Use mimetypes as fallback if python-magic is not available
        self.use_magic = False
        try:
            import magic
            self.mime_detector = magic.Magic(mime=True)
            self.use_magic = True
        except ImportError:
            logger.warning("python-magic not available, using mimetypes as fallback")
            
    def _get_mime_type(self, file_path: Path) -> str:
        """Get MIME type using available method"""
        if self.use_magic:
            return self.mime_detector.from_file(str(file_path))
        else:
            # Fallback to mimetypes
            mime_type, _ = mimetypes.guess_type(str(file_path))
            return mime_type or "application/octet-stream"
        
    async def get_file_info(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Extract comprehensive file information"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Basic file info
        stat = file_path.stat()
        file_info = {
            "name": file_path.name,
            "path": str(file_path.absolute()),
            "size": stat.st_size,
            "size_mb": round(stat.st_size / (1024 * 1024), 2),
            "created_at": stat.st_ctime,
            "modified_at": stat.st_mtime,
            "extension": file_path.suffix.lower(),
            "mime_type": self._get_mime_type(file_path)
        }
        
        # File type classification
        file_info["category"] = self._classify_file_type(file_info["mime_type"], file_info["extension"])
        
        # Generate file hash for duplicate detection
        file_info["hash"] = await self._generate_file_hash(file_path)
        
        # Type-specific metadata
        try:
            if file_info["category"] == "image":
                file_info["metadata"] = await self._extract_image_metadata(file_path)
            elif file_info["category"] == "video":
                file_info["metadata"] = await self._extract_video_metadata(file_path)
            elif file_info["category"] == "audio":
                file_info["metadata"] = await self._extract_audio_metadata(file_path)
            elif file_info["category"] == "document":
                file_info["metadata"] = await self._extract_document_metadata(file_path)
            else:
                file_info["metadata"] = {}
        except Exception as e:
            logger.warning(f"Failed to extract metadata for {file_path}: {e}")
            file_info["metadata"] = {}
        
        return file_info
    
    def _classify_file_type(self, mime_type: str, extension: str) -> str:
        """Classify file into broad categories"""
        if mime_type.startswith("image/"):
            return "image"
        elif mime_type.startswith("video/"):
            return "video"
        elif mime_type.startswith("audio/"):
            return "audio"
        elif mime_type in ["application/pdf", "application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            return "document"
        elif mime_type in ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "text/csv"]:
            return "spreadsheet"
        elif mime_type.startswith("text/") or extension in [".md", ".txt", ".py", ".js", ".html", ".css"]:
            return "text"
        elif mime_type in ["application/zip", "application/x-rar-compressed", "application/x-7z-compressed"]:
            return "archive"
        else:
            return "other"
    
    async def _generate_file_hash(self, file_path: Path) -> str:
        """Generate SHA-256 hash for file"""
        hash_sha256 = hashlib.sha256()
        
        async with aiofiles.open(file_path, 'rb') as f:
            while chunk := await f.read(8192):
                hash_sha256.update(chunk)
                
        return hash_sha256.hexdigest()
    
    async def _extract_image_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract image-specific metadata"""
        metadata = {}
        
        if not PIL_AVAILABLE:
            return metadata
        
        try:
            with Image.open(file_path) as img:
                metadata.update({
                    "dimensions": img.size,
                    "width": img.size[0],
                    "height": img.size[1],
                    "mode": img.mode,
                    "format": img.format,
                    "has_transparency": img.mode in ('RGBA', 'LA') or 'transparency' in img.info
                })
                
                # EXIF data
                if hasattr(img, '_getexif') and img._getexif():
                    exif = img._getexif()
                    if exif:
                        metadata["exif"] = {k: str(v) for k, v in exif.items() if isinstance(v, (str, int, float))}
                
                # Perceptual hash for similarity detection
                if IMAGEHASH_AVAILABLE:
                    phash = imagehash.phash(img)
                    metadata["perceptual_hash"] = str(phash)
                
        except Exception as e:
            logger.warning(f"Failed to extract image metadata: {e}")
            
        return metadata
    
    async def _extract_video_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract video-specific metadata"""
        metadata = {}
        
        if not CV2_AVAILABLE:
            return metadata
        
        try:
            cap = cv2.VideoCapture(str(file_path))
            
            if cap.isOpened():
                metadata.update({
                    "width": int(cap.get(cv2.CAP_PROP_FRAME_WIDTH)),
                    "height": int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT)),
                    "fps": cap.get(cv2.CAP_PROP_FPS),
                    "frame_count": int(cap.get(cv2.CAP_PROP_FRAME_COUNT)),
                    "duration_seconds": int(cap.get(cv2.CAP_PROP_FRAME_COUNT)) / cap.get(cv2.CAP_PROP_FPS) if cap.get(cv2.CAP_PROP_FPS) > 0 else 0
                })
                
            cap.release()
            
        except Exception as e:
            logger.warning(f"Failed to extract video metadata: {e}")
            
        return metadata
    
    async def _extract_audio_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract audio-specific metadata"""
        metadata = {}
        
        try:
            # Try to import mutagen if available
            from mutagen import File as MutagenFile
            
            audio_file = MutagenFile(file_path)
            
            if audio_file:
                info = audio_file.info
                metadata.update({
                    "duration_seconds": getattr(info, 'length', 0),
                    "bitrate": getattr(info, 'bitrate', 0),
                    "channels": getattr(info, 'channels', 0),
                    "sample_rate": getattr(info, 'sample_rate', 0)
                })
                
                # Tags
                if audio_file.tags:
                    tags = {}
                    for key, value in audio_file.tags.items():
                        if isinstance(value, list):
                            tags[key] = [str(v) for v in value]
                        else:
                            tags[key] = str(value)
                    metadata["tags"] = tags
        
        except ImportError:
            logger.warning("mutagen not available, skipping audio metadata extraction")
        except Exception as e:
            logger.warning(f"Failed to extract audio metadata: {e}")
            
        return metadata
    
    async def _extract_document_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract document-specific metadata"""
        metadata = {}
        extension = file_path.suffix.lower()
        
        try:
            if extension == ".pdf" and PYPDF_AVAILABLE:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata.update({
                        "pages": len(pdf_reader.pages),
                        "encrypted": pdf_reader.is_encrypted
                    })
                    
                    if pdf_reader.metadata:
                        pdf_metadata = {k[1:] if k.startswith('/') else k: str(v) for k, v in pdf_reader.metadata.items()}
                        metadata["pdf_info"] = pdf_metadata
                        
            elif extension == ".docx" and DOCX_AVAILABLE:
                doc = Document(file_path)
                metadata.update({
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "images": len(doc.inline_shapes)
                })
                
                if doc.core_properties:
                    core_props = {}
                    for prop in ['author', 'category', 'comments', 'content_status', 'created', 'identifier',
                                'keywords', 'language', 'last_modified_by', 'last_printed', 'modified',
                                'revision', 'subject', 'title', 'version']:
                        value = getattr(doc.core_properties, prop, None)
                        if value:
                            core_props[prop] = str(value)
                    metadata["core_properties"] = core_props
                    
            elif extension in [".xlsx", ".xls"] and OPENPYXL_AVAILABLE:
                workbook = openpyxl.load_workbook(file_path, read_only=True)
                metadata.update({
                    "worksheets": len(workbook.sheetnames),
                    "sheet_names": workbook.sheetnames
                })
                
                # Count total rows and columns
                total_rows = 0
                total_cols = 0
                for sheet in workbook.worksheets:
                    total_rows += sheet.max_row
                    total_cols = max(total_cols, sheet.max_column)
                    
                metadata.update({
                    "total_rows": total_rows,
                    "max_columns": total_cols
                })
                
        except Exception as e:
            logger.warning(f"Failed to extract document metadata: {e}")
            
        return metadata
    
    async def read_text_content(self, file_path: Union[str, Path], max_length: Optional[int] = None) -> str:
        """Extract text content from various file types"""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        content = ""
        
        try:
            if extension == ".pdf":
                content = await self._extract_pdf_text(file_path)
            elif extension == ".docx":
                content = await self._extract_docx_text(file_path)
            elif extension in [".txt", ".md", ".py", ".js", ".html", ".css", ".json", ".xml"]:
                async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = await f.read()
            elif extension in [".xlsx", ".xls"]:
                content = await self._extract_excel_text(file_path)
            else:
                logger.warning(f"Text extraction not supported for {extension} files")
                
            if max_length and len(content) > max_length:
                content = content[:max_length] + "...[truncated]"
                
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            
        return content
    
    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        text = ""
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
                
        return text.strip()
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from Word document"""
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
                    
        return text.strip()
    
    async def _extract_excel_text(self, file_path: Path) -> str:
        """Extract text from Excel file"""
        workbook = openpyxl.load_workbook(file_path, read_only=True)
        text = ""
        
        for sheet in workbook.worksheets:
            text += f"\n=== Sheet: {sheet.title} ===\n"
            
            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                if row_text.strip():
                    text += row_text + "\n"
                    
        return text.strip()
    
    async def detect_duplicates(self, file_paths: List[Path], similarity_threshold: float = 0.9) -> List[List[Path]]:
        """Detect duplicate and similar files"""
        duplicates = []
        processed = set()
        
        for i, file_path in enumerate(file_paths):
            if file_path in processed:
                continue
                
            similar_files = [file_path]
            
            try:
                file_info = await self.get_file_info(file_path)
                
                # Check exact duplicates by hash
                for j, other_path in enumerate(file_paths[i+1:], i+1):
                    if other_path in processed:
                        continue
                        
                    other_info = await self.get_file_info(other_path)
                    
                    # Exact duplicate (same hash)
                    if file_info["hash"] == other_info["hash"]:
                        similar_files.append(other_path)
                        processed.add(other_path)
                    
                    # Image similarity (perceptual hash)
                    elif (IMAGEHASH_AVAILABLE and 
                          file_info["category"] == "image" and other_info["category"] == "image" and
                          "perceptual_hash" in file_info.get("metadata", {}) and
                          "perceptual_hash" in other_info.get("metadata", {})):
                        
                        hash1 = imagehash.hex_to_hash(file_info["metadata"]["perceptual_hash"])
                        hash2 = imagehash.hex_to_hash(other_info["metadata"]["perceptual_hash"])
                        
                        similarity = 1 - (hash1 - hash2) / len(hash1.hash) ** 2
                        
                        if similarity >= similarity_threshold:
                            similar_files.append(other_path)
                            processed.add(other_path)
                            
            except Exception as e:
                logger.warning(f"Failed to process {file_path} for duplicate detection: {e}")
                
            if len(similar_files) > 1:
                duplicates.append(similar_files)
                
            processed.add(file_path)
            
        return duplicates


# Global file processor instance
file_processor = FileProcessor()
